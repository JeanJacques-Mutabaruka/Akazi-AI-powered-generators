"""
AKAZI Job Description Generator V4 — Hérite de BaseGenerator
Pattern Template Method identique aux générateurs CV.

CHANGEMENTS V4 vs V3 :
  ✅ Hérite BaseGenerator  → supprime load_data, save_document, generate dupliqués
  ✅ setup_document()      → surcharge légère (marges 2.54 cm, désactive spacing)
  ✅ _apply_header_footer()→ délégué à BaseGenerator (preset YAML ou fallback AKAZI)
  ✅ validate_data()       → signature harmonisée (self.data déjà chargé)
  ✅ generate_content()    → remplace generate() comme point d'entrée contenu
  ✅ _make_run()           → remplace les blocs run.font.name/size/color.rgb
  ✅ _apply_para_format()  → remplace _apply_paragraph_format() (même API BaseGenerator)
  ✅ _add_bullet_item()    → hérité de BaseGenerator (remplace _add_bullet_level1/2)
  ✅ _add_paragraph()      → conservé (spécifique JD, pas dans BaseGenerator)
  ✅ _create_section_title → remplace _create_section_title (utilise _make_run)
  ✅ Styles bullet         → délégués à AkaziStyleManager.setup_all_styles() (déjà fait)
  ✅ Suppression méthodes  → _create_akazi_bullet_level1_style, _make_style_visible_in_gallery
                             _disable_contextual_spacing_for_style (dans BaseGenerator)
"""

from pathlib import Path
from typing import Dict, Any, Optional, List

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from generators.base_generator import BaseGenerator


class AkaziJobDescGenerator(BaseGenerator):
    """
    Générateur Fiche de Poste au format AKAZI (EN/FR).
    Structure JSON attendue : document_metadata, mission_title, budget, description.sections
    """

    DEFAULT_FONT = "Century Gothic"
    DEFAULT_SIZE = 10

    # ── Paramètres de formatage standardisés ─────────────────────────────────
    LINE_SPACING_MULTIPLE       = 1.15
    SECTION_TITLE_SPACE_BEFORE  = 20
    SECTION_TITLE_SPACE_AFTER   = 6
    BULLET_SPACE_BEFORE         = 3
    BULLET_SPACE_AFTER          = 3
    BULLET_L1_LEFT_INDENT       = 1.0   # cm
    BULLET_L1_FIRST_LINE        = -0.5  # cm (hanging)
    BULLET_L2_LEFT_INDENT       = 1.5   # cm
    BULLET_L2_FIRST_LINE        = -0.5  # cm (hanging)

    def __init__(self, input_file: Path, output_file: Path, lang: str = 'fr', **kwargs):
        super().__init__(input_file, output_file, lang=lang, **kwargs)

    # =========================================================================
    # SETUP DOCUMENT — surcharge pour marges JD (2.54 cm) + désactivations
    # =========================================================================

    def setup_document(self):
        """
        Surcharge : marges AKAZI JD (2.54 cm), désactive contextual spacing.
        NE configure PAS le header/footer — délégué à _apply_header_footer().
        """
        super().setup_document()  # styles AKAZI + police Normal

        # Marges Job Description (2.54 cm vs 1.27 cm pour les CV)
        self._set_margins(Cm(2.54))

        # Désactiver contextual spacing (hérité de BaseGenerator)
        for style_name in ('Normal', 'List Bullet', 'List Bullet 2'):
            self._disable_contextual_spacing(style_name)

    # =========================================================================
    # VALIDATE DATA
    # =========================================================================

    def validate_data(self) -> bool:
        metadata    = self.data.get('document_metadata', {})
        doc_type    = metadata.get('document_type')
        format_code = metadata.get('format_code')

        if doc_type != 'job_description':
            raise ValueError(f"document_type incorrect : '{doc_type}' (attendu : 'job_description')")
        if not format_code or 'AKAZI' not in format_code:
            raise ValueError(f"format_code incorrect : '{format_code}'")
        return True

    # =========================================================================
    # GENERATE CONTENT — pipeline JD
    # =========================================================================

    def generate_content(self):
        """Pipeline contenu AKAZI JD : titre → budget → intro → sections → footer"""
        self.generate_header()
        self.generate_budget()

        description = self.data.get('description', {})

        # Intro optionnel
        intro = description.get('intro', '')
        if intro:
            self._add_paragraph(intro, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

        # Sections (nouveau format)
        if 'sections' in description:
            self._generate_sections(description['sections'])

        # Footer optionnel
        footer_text = description.get('footer', '')
        if footer_text:
            self._add_paragraph(footer_text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # =========================================================================
    # HEADER — titre de la fiche de poste
    # =========================================================================

    def generate_header(self):
        """
        Titre de la fiche de poste.
        Rouge, gras, 14pt, centré.
        """
        # Compatibilité ancien format
        if 'mission_title' in self.data:
            title = self.data.get('mission_title', 'FICHE DE POSTE')
        else:
            title = self.data.get('job_information', {}).get('job_title', 'FICHE DE POSTE')

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_para_format(para, space_before=0, space_after=0,
                                 line_spacing=self.LINE_SPACING_MULTIPLE)
        self._make_run(para, title.upper(), size=14, color=self.COLOR_BLACK, bold=True)

    # =========================================================================
    # BUDGET — ligne rouge unique "BUDGET : [valeur]"
    # =========================================================================

    def generate_budget(self):
        """
        BUDGET : [valeur]
        Rouge, gras, 12pt, gauche. Une seule ligne.
        """
        budget_data = self.data.get('budget', {})
        if not budget_data:
            return

        if 'text' in budget_data:
            budget_text = budget_data['text']
        else:
            # Compatibilité ancien format
            budget_text = budget_data.get('daily_rate', '') or budget_data.get('total_budget', '')

        if not budget_text:
            return

        para = self.doc.add_paragraph()
        self._apply_para_format(
            para,
            space_before=self.SECTION_TITLE_SPACE_BEFORE,
            space_after=self.SECTION_TITLE_SPACE_AFTER,
            line_spacing=self.LINE_SPACING_MULTIPLE
        )
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._make_run(para, f"BUDGET : {budget_text}", size=12, color=self.COLOR_RED, bold=True)

    # =========================================================================
    # SECTIONS — titre + contenu (bullets ou texte)
    # =========================================================================

    def _generate_sections(self, sections: List[Dict]):
        """Génère toutes les sections de la fiche de poste."""
        for section in sections:
            title   = section.get('title', '')
            content = section.get('content', '')

            if title:
                self._create_section_title(title)

            # Détecter section compétences pour mise en gras
            is_skills = any(kw in title.upper() for kw in [
                'COMPÉTENCES', 'QUALIFICATIONS', 'SKILLS', 'REQUIRED SKILLS', 'COMPETENCIES'
            ])

            if isinstance(content, str):
                bold_terms = section.get('formatting', {}).get('bold_terms', [])
                if bold_terms:
                    para = self.doc.add_paragraph()
                    self._apply_para_format(para, space_before=3, space_after=3,
                                             line_spacing=self.LINE_SPACING_MULTIPLE)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self._add_text_with_bold_terms(para, content, bold_terms)
                else:
                    self._add_paragraph(content)

            elif isinstance(content, list):
                self._process_content_list(content, is_skills_section=is_skills)

    def _create_section_title(self, text: str, is_budget: bool = False):
        """
        Titre de section : 12pt, gras, NOIR (ROUGE seulement pour budget).
        Utilise _make_run() de BaseGenerator.
        """
        color = self.COLOR_RED if is_budget else self.COLOR_BLACK

        para = self.doc.add_paragraph(style='Heading 3')
        self._apply_para_format(
            para,
            space_before=self.SECTION_TITLE_SPACE_BEFORE,
            space_after=self.SECTION_TITLE_SPACE_AFTER,
            line_spacing=self.LINE_SPACING_MULTIPLE
        )
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._make_run(para, text.upper(), size=12, color=color, bold=True)

    def _process_content_list(self, content: List, is_skills_section: bool = False):
        """
        Traite une liste de bullets (niveau 1 + sous-items niveau 2).
        Délègue à _add_bullet_item() de BaseGenerator pour les bullets.
        """
        for item in content:
            if isinstance(item, str):
                self._add_jd_bullet_l1(item, make_bold=is_skills_section)

            elif isinstance(item, dict):
                text      = item.get('text', '')
                inferred  = item.get('inferred', False)
                sub_items = item.get('sub_items', [])
                color     = self.COLOR_ORANGE if inferred else self.COLOR_BLACK

                self._add_jd_bullet_l1(text, make_bold=is_skills_section, color=color)

                for sub in sub_items:
                    if isinstance(sub, str):
                        self._add_jd_bullet_l2(sub)
                    else:
                        sub_text     = sub.get('text', '')
                        sub_inferred = sub.get('inferred', False)
                        self._add_jd_bullet_l2(
                            sub_text,
                            color=self.COLOR_ORANGE if sub_inferred else self.COLOR_BLACK
                        )

    # =========================================================================
    # HELPERS BULLET JD — spécificités du format Job Description
    # (indentations et line_spacing différents de _add_bullet_item() BaseGenerator)
    # =========================================================================

    def _add_jd_bullet_l1(self, text: str, make_bold: bool = False, color=None):
        """
        Bullet niveau 1 JD.
        Utilise _add_bullet_item() de BaseGenerator avec paramètres JD spécifiques.
        make_bold : True pour la section Compétences (tout le texte en gras).
        """
        if color is None:
            color = self.COLOR_BLACK

        para = self._add_bullet_item(
            self.doc, text,
            color       = color,
            left_indent = self.BULLET_L1_LEFT_INDENT,
            hanging     = abs(self.BULLET_L1_FIRST_LINE),
            size        = self.DEFAULT_SIZE,
            level       = 0,
        )

        # Cas section compétences : passer tout le run en gras
        if make_bold and para is not None:
            for run in para.runs:
                run.bold = True

        # Line spacing spécifique JD
        if para is not None:
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = self.LINE_SPACING_MULTIPLE

    def _add_jd_bullet_l2(self, text: str, color=None):
        """
        Bullet niveau 2 JD (sous-item).
        Utilise _add_bullet_item() de BaseGenerator avec paramètres L2.
        """
        if color is None:
            color = self.COLOR_BLACK

        para = self._add_bullet_item(
            self.doc, text,
            color       = color,
            left_indent = self.BULLET_L2_LEFT_INDENT,
            hanging     = abs(self.BULLET_L2_FIRST_LINE),
            size        = self.DEFAULT_SIZE,
            level       = 1,
        )

        if para is not None:
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = self.LINE_SPACING_MULTIPLE

    # =========================================================================
    # HELPERS TEXTE — spécifiques à la JD, pas dans BaseGenerator
    # =========================================================================

    def _add_paragraph(self, text: str, bold: bool = False,
                        color=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """Paragraphe normal avec formatage standard JD."""
        if color is None:
            color = self.COLOR_BLACK

        para = self.doc.add_paragraph()
        self._apply_para_format(para, space_before=3, space_after=3,
                                 line_spacing=self.LINE_SPACING_MULTIPLE)
        para.alignment = alignment
        self._make_run(para, text, size=self.DEFAULT_SIZE, color=color, bold=bold)

    def _add_text_with_bold_terms(self, para, text: str, bold_terms: List[str]):
        """
        Ajoute du texte dans un paragraphe existant en mettant
        certains termes en gras.
        """
        remaining = text
        while remaining:
            next_term = None
            next_pos  = len(remaining)

            for term in bold_terms:
                pos = remaining.find(term)
                if pos != -1 and pos < next_pos:
                    next_pos  = pos
                    next_term = term

            if next_term:
                if next_pos > 0:
                    self._make_run(para, remaining[:next_pos],
                                   size=self.DEFAULT_SIZE, color=self.COLOR_BLACK)
                self._make_run(para, next_term,
                               size=self.DEFAULT_SIZE, color=self.COLOR_BLACK, bold=True)
                remaining = remaining[next_pos + len(next_term):]
            else:
                self._make_run(para, remaining,
                               size=self.DEFAULT_SIZE, color=self.COLOR_BLACK)
                remaining = ""
