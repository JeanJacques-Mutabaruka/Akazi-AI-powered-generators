"""
MC2I CV Generator V5
Hérite de BaseGenerator V3.

CHANGEMENTS V5 :
  ✅ Point 3 : generate_header() retiré du pipeline generate_content()
     (les 3 lignes CONSULTANT DATA / rôle / années n'apparaissent plus dans
     le corps du document — elles restent dans le header Word via le preset YAML)
  ✅ Point 4 : generate_experience_summary() ajoutée et insérée dans le pipeline
     entre generate_languages() et generate_education() — section SYNTHÈSE DES
     EXPÉRIENCES manquante dans V4
  ✅ Point 5 : _add_mc2i_section_title() utilise désormais COLOR_COMPANY
     (RGB 221,0,97) au lieu de COLOR_TEXT pour les titres de sections
  ✅ V4 : _apply_header_footer(), _build_hf_variables(), preset YAML dynamique
"""

from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generators.base_generator import BaseGenerator
from utils.logger import logger


# Preset MC2I par défaut (relatif à la racine du projet)
_DEFAULT_MC2I_PRESET = "hf_presets/combined/mc2i_cv_complet.yaml"


class MC2ICVGenerator(BaseGenerator):
    """
    Générateur CV au format MC2I (Dossier de Compétences).
    Structure JSON attendue : document_metadata, introduction, professional_experiences
    """

    DEFAULT_FONT = "Lato"
    DEFAULT_SIZE = 10

    def __init__(self, input_file: Path, output_file: Path, **kwargs):
        # Si aucun preset explicite fourni → on utilise le preset MC2I par défaut
        if "hf_preset" not in kwargs or kwargs["hf_preset"] is None:
            kwargs["hf_preset"] = _DEFAULT_MC2I_PRESET
        super().__init__(input_file, output_file, **kwargs)

    # =========================================================================
    # SETUP DOCUMENT
    # =========================================================================

    def setup_document(self):
        """
        Marges MC2I (2.54 cm), police Lato 10pt, texte justifié.
        NE configure PAS le header/footer — délégué à _apply_header_footer().
        """
        self.doc = Document()

        try:
            from utils.akazi_styles import AkaziStyleManager
            AkaziStyleManager(self.doc).setup_all_styles()
        except Exception as e:
            logger.warning("style_manager_skipped", reason=str(e))

        normal = self.doc.styles["Normal"]
        normal.font.name      = self.DEFAULT_FONT
        normal.font.size      = Pt(self.DEFAULT_SIZE)
        normal.font.color.rgb = self.COLOR_TEXT

        pf = normal.paragraph_format
        pf.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.15

        self._set_margins(Cm(2.54))

    # =========================================================================
    # HEADER / FOOTER — preset YAML + variables dynamiques  (V4)
    # =========================================================================

    def _apply_header_footer(self):
        """
        Applique le preset MC2I avec substitution des variables dynamiques.

        Stratégie unique (V4) :
          1. Résoudre le chemin du preset (hf_preset ou défaut MC2I)
          2. Extraire les variables consultant depuis self.data
          3. Appeler HeaderFooterEngine.apply_yaml_with_vars()

        Plus de branchement vers MC2ILayoutManager — le preset YAML
        est la seule source de vérité pour le layout MC2I.
        """
        # hf_preset == "none" → pas de header/footer du tout
        if self.hf_preset == "none":
            logger.debug("hf_skipped", reason="hf_preset='none'")
            return

        preset_path = self._resolve_preset_path()
        if preset_path is None:
            logger.warning(
                "mc2i_hf_preset_not_found",
                preset=str(self.hf_preset),
                reason="Fichier introuvable — header/footer ignoré"
            )
            return

        variables = self._build_hf_variables()

        try:
            from docx_header_engine.engine import HeaderFooterEngine
            engine = HeaderFooterEngine(self.doc)
            engine.apply_yaml_with_vars(str(preset_path), variables=variables)
            logger.info(
                "mc2i_hf_preset_applied",
                preset=str(preset_path),
                variables=variables,
            )
        except Exception as e:
            import traceback as _tb
            _detail = _tb.format_exc()
            logger.error("mc2i_hf_preset_failed", error=str(e), detail=_detail)
            self._hf_error = (
                f"{type(e).__name__}: {e}\n"
                f"Preset : {preset_path}\n"
                f"Variables : {variables}\n"
                f"{_detail}"
            )

    def _resolve_preset_path(self) -> Optional[Path]:
        """
        Résout le chemin absolu du preset YAML.
        Cherche d'abord relatif à la racine du projet, puis en absolu.
        """
        candidate = Path(self.hf_preset)

        # Chemin absolu fourni directement
        if candidate.is_absolute():
            return candidate if candidate.exists() else None

        # Chemin relatif → résoudre depuis la racine du projet
        project_root = Path(__file__).resolve().parent.parent
        resolved = project_root / candidate
        return resolved if resolved.exists() else None

    def _build_hf_variables(self) -> Dict[str, str]:
        """
        Extrait les variables dynamiques depuis self.data
        pour alimenter les {{placeholders}} du preset YAML.

        Variables produites :
            -- Consultant --
            consultant_initials  →  "NKA"
            consultant_name      →  "JEAN DUPONT"
            consultant_email     →  "j.dupont@mc2i.fr"
            agency_address       →  "51 Rue François 1er – 75008 PARIS…"
            title                →  "Tech Lead Data"
            years                →  "12"

            -- Document (tout document_metadata) --
            format_code          →  "MC2I_V1"
            language_iso         →  "FRA"
            document_type        →  "cv"
            generated_at         →  "2026-04-01T12:00:00Z"

            -- Calculés --
            main_domain          →  "Big Data et Cloud Computing"

        Returns:
            Dict {nom: valeur_string} — toutes les valeurs sont des strings.
        """
        metadata    = self.data.get("document_metadata", {})
        exp_summary = self.data.get("experience_summary", [])

        # ── Initiales / nom / email / adresse ─────────────────────────────
        initials         = metadata.get("consultant_initials", "")
        consultant_name  = metadata.get("consultant_name", "")
        consultant_email = metadata.get("consultant_email", "")
        agency_address   = metadata.get(
            "agency_address",
            "51 Rue François 1er – 75008 PARIS - 01.44.43.01.00 – www.mc2i.fr"
        )

        # ── Titre de mission : priorité experience_summary, fallback pro_exp ─
        title = ""
        if exp_summary and isinstance(exp_summary[0], dict):
            title = exp_summary[0].get("title", "")
        if not title:
            pro_exp = self.data.get("professional_experiences", [])
            if pro_exp and isinstance(pro_exp[0], dict):
                title = pro_exp[0].get("title", "Consultant")

        # ── Années calculées ───────────────────────────────────────────────
        total_months = sum(
            e.get("duration_months", 0)
            for e in exp_summary
            if isinstance(e, dict)
        )
        years = str(total_months // 12)

        # ── main_domain depuis introduction.conclusion ─────────────────────
        conclusion  = self.data.get("introduction", {}).get("conclusion", {})
        main_domain = ""
        if isinstance(conclusion, dict):
            main_domain = conclusion.get("main_domain", "")

        # ── Tout document_metadata en string (clés plates) ─────────────────
        meta_vars = {
            k: str(v) if v is not None else ""
            for k, v in metadata.items()
            if isinstance(v, (str, int, float, bool)) or v is None
        }

        # ── Assembler le dict final (meta_vars en base, spécifiques par-dessus) ─
        variables = {**meta_vars}
        variables.update({
            "consultant_initials":    initials,
            "consultant_name":        consultant_name,
            "consultant_email":       consultant_email,
            "agency_address":         agency_address,
            "years_experience":       years,           # ← renommé (était "years")
            "main_domain_expertise":  main_domain,     # ← renommé (était "main_domain")
        })

        return variables


    # =========================================================================
    # VALIDATE DATA  (inchangé)
    # =========================================================================

    def validate_data(self) -> bool:
        required = ["document_metadata", "introduction", "professional_experiences"]
        missing  = [k for k in required if k not in self.data]
        if missing:
            raise ValueError(f"Clés manquantes dans les données MC2I CV : {missing}")
        fmt = self.data["document_metadata"].get("format_code")
        if fmt != "MC2I_V1":
            raise ValueError(f"format_code incorrect : '{fmt}' (attendu : 'MC2I_V1')")
        return True

    # =========================================================================
    # GENERATE CONTENT  (inchangé)
    # =========================================================================

    def generate_content(self):
        """
        Pipeline contenu MC2I :
          intro → langues → synthèse expériences → formation → expertise → expériences

        Note V5 :
          - generate_header() retiré (Point 3) : les lignes CONSULTANT DATA /
            rôle / années sont désormais uniquement dans le header Word (preset YAML).
          - generate_experience_summary() ajoutée (Point 4) : section SYNTHÈSE
            DES EXPÉRIENCES insérée entre langues et formation.
        """
        self.generate_introduction()
        self._add_horizontal_separator()
        self.generate_languages()
        self._add_horizontal_separator()
        self.generate_experience_summary()
        self._add_horizontal_separator()
        self.generate_education()
        self._add_horizontal_separator()
        self.generate_expertise()
        self._add_horizontal_separator()
        self.generate_experiences()

    # =========================================================================
    # HEADER CONTENU  (inchangé)
    # =========================================================================

    def generate_header(self):
        """
        GAP 4 — Hiérarchie corrigée pour correspondre au template MC2I cible :

        Ligne 1 (H1) : "CONSULTANT DATA"
                       → Lato 16pt, gras, COLOR_TEXT (gris #575856)
        Ligne 2 (H2) : Titre du rôle principal (ex: "Tech Lead Data")
                       → Lato 14pt, gras, COLOR_COMPANY (rose #DD0061)
        Ligne 3 (H3) : Années d'expérience (ex: "12 années d'expérience")
                       → Lato 10pt, normal, COLOR_TEXT
        """
        exp_sum = self.data.get("experience_summary", [])

        # ── Titre du rôle principal ────────────────────────────────────────
        # Priorité : experience_summary[0]['title'] → professional_experiences[0]['title']
        role_title = ""
        if exp_sum and isinstance(exp_sum[0], dict):
            role_title = exp_sum[0].get("title", "")
        if not role_title:
            pro_exp = self.data.get("professional_experiences", [])
            if pro_exp and isinstance(pro_exp[0], dict):
                role_title = pro_exp[0].get("title", "Consultant")
        if not role_title:
            role_title = "Consultant"

        # ── Années totales ─────────────────────────────────────────────────
        total_months = sum(
            e.get("duration_months", 0)
            for e in exp_sum
            if isinstance(e, dict)
        )
        years = total_months // 12

        # ── H1 : "CONSULTANT DATA" — gris 16pt bold ────────────────────────
        p_h1 = self.doc.add_paragraph()
        p_h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_para_format(p_h1, space_before=0, space_after=2)
        self._make_run(p_h1, "CONSULTANT DATA",
                       font="Lato", size=16, color=self.COLOR_TEXT, bold=True)

        # ── H2 : Titre du rôle — rose 14pt bold ───────────────────────────
        p_h2 = self.doc.add_paragraph()
        p_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_para_format(p_h2, space_before=0, space_after=2)
        self._make_run(p_h2, role_title,
                       font="Lato", size=14, color=self.COLOR_COMPANY, bold=True)

        # ── H3 : Années d'expérience — gris 10pt normal ───────────────────
        p_h3 = self.doc.add_paragraph()
        p_h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_para_format(p_h3, space_before=0, space_after=4)
        self._make_run(p_h3, f"{years} années d'expérience",
                       font="Lato", size=10, color=self.COLOR_TEXT)

    # =========================================================================
    # INTRODUCTION  (inchangé)
    # =========================================================================

    def generate_introduction(self):
        intro = self.data.get("introduction", {})
        for key in ["experience_summary", "technical_skills_summary", "functional_skills_summary"]:
            text = intro.get(key, "")
            if text:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._make_run(p, text, font="Lato", size=10, color=self.COLOR_TEXT)

        conclusion = intro.get("conclusion", {})
        conclusion_text = (
            conclusion.get("text", "") if isinstance(conclusion, dict) else str(conclusion)
        )
        if conclusion_text:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._make_run(p, conclusion_text, font="Lato", size=10, color=self.COLOR_TEXT)

    # =========================================================================
    # LANGUES  (inchangé)
    # =========================================================================

    def generate_languages(self):
        self._add_mc2i_section_title("LANGUES")
        for lang in self.data.get("languages", []):
            self._add_bullet_item(
                self.doc,
                f"{lang.get('language', '')} : {lang.get('level', '')}",
                color=self.COLOR_TEXT, font="Lato", size=10
            )

    # =========================================================================
    # SYNTHÈSE DES EXPÉRIENCES  (Point 4 — ajouté V5)
    # =========================================================================

    def generate_experience_summary(self):
        """
        Section SYNTHÈSE DES EXPÉRIENCES :
        Liste toutes les expériences avec entreprise, titre et durée en mois.
        Source : data['experience_summary'] (liste compacte générée par le prompt).
        Format : • Entreprise — Titre — N mois
        """
        exp_summary = self.data.get("experience_summary", [])
        if not exp_summary:
            return

        label = "SYNTHÈSE DES EXPÉRIENCES"
        self._add_mc2i_section_title(label)

        for exp in exp_summary:
            if not isinstance(exp, dict):
                continue
            company  = exp.get("company", "")
            title    = exp.get("title", "")
            months   = exp.get("duration_months", 0)

            # Construire la ligne : "Entreprise — Titre — N mois"
            parts = [p for p in [company, title] if p]
            line  = " — ".join(parts)
            if months:
                line = f"{line} — {months} mois"

            self._add_bullet_item(
                self.doc, line,
                color=self.COLOR_TEXT, font="Lato", size=10
            )

    # =========================================================================
    # FORMATION  (inchangé)
    # =========================================================================

    def generate_education(self):
        self._add_mc2i_section_title("FORMATION")
        for edu in self.data.get("education", []):
            self._add_bullet_item(
                self.doc,
                f"{edu.get('year', '')} : {edu.get('degree', '')} - {edu.get('institution', '')}",
                color=self.COLOR_TEXT, font="Lato", size=10
            )

    # =========================================================================
    # EXPERTISE  (inchangé)
    # =========================================================================

    def generate_expertise(self):
        self._add_mc2i_section_title("EXPERTISES, OUTILS ET TECHNOLOGIES")
        expertise = self.data.get("expertise", {})
        for exp in expertise.get("expertises", []):
            self._add_bullet_item(self.doc, exp, color=self.COLOR_TEXT, font="Lato", size=10)
        masteries = expertise.get("masteries", [])
        if masteries:
            self._add_bullet_item(self.doc, ", ".join(masteries),
                                   color=self.COLOR_TEXT, font="Lato", size=10)

    # =========================================================================
    # EXPÉRIENCES  (inchangé)
    # =========================================================================

    def generate_experiences(self):
        self._add_mc2i_section_title("EXPÉRIENCES PROFESSIONNELLES")
        experiences = self.data.get("professional_experiences", [])
        for idx, exp in enumerate(experiences):
            self._generate_single_experience(exp)
            if idx < len(experiences) - 1:
                self._add_horizontal_separator()

    def _generate_single_experience(self, exp: Dict):
        p_company = self.doc.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._make_run(p_company, exp.get("company", "N/A").upper(),
                       font="Lato", size=14, color=self.COLOR_COMPANY,
                       bold=True, small_caps=True)

        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._make_run(p_title, exp.get("title", "N/A"),
                       font="Lato", size=14, color=self.COLOR_MISSION,
                       bold=True, small_caps=True)

        period = exp.get("period", {}).get("formatted", "N/A")
        p_period = self.doc.add_paragraph()
        p_period.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._make_run(p_period, period, font="Lato", size=10,
                       color=self.COLOR_MISSION, small_caps=True)

        context = exp.get("context", "")
        if context:
            p_ctx = self.doc.add_paragraph()
            p_ctx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._make_run(p_ctx, context, font="Lato", size=10, color=self.COLOR_TEXT)

        p_act_label = self.doc.add_paragraph()
        self._make_run(p_act_label, "Activités :",
                       font="Lato", size=10, color=self.COLOR_TEXT, bold=True)
        for activity in exp.get("activities", []):
            self._add_bullet_item(self.doc, activity,
                                   color=self.COLOR_TEXT, font="Lato", size=10)

        p_dom_label = self.doc.add_paragraph()
        self._make_run(p_dom_label, "Domaines :",
                       font="Lato", size=10, color=self.COLOR_TEXT, bold=True)
        for dom in exp.get("functional_domains", []):
            domain_text = dom.get("domain", "") if isinstance(dom, dict) else str(dom)
            self._add_bullet_item(self.doc, domain_text,
                                   color=self.COLOR_TEXT, font="Lato", size=10)

        p_tech_label = self.doc.add_paragraph()
        self._make_run(p_tech_label, "Environnement technique :",
                       font="Lato", size=10, color=self.COLOR_TEXT, bold=True)

        tech_env   = exp.get("technical_environment", [])
        tech_texts = [
            t.get("technologies", "") if isinstance(t, dict) else str(t)
            for t in tech_env
        ]
        if tech_texts:
            p_tech = self.doc.add_paragraph()
            p_tech.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._make_run(p_tech, ", ".join(filter(None, tech_texts)),
                           font="Lato", size=10, color=self.COLOR_TEXT)

    # =========================================================================
    # HELPER SPÉCIFIQUE MC2I  (inchangé)
    # =========================================================================

    def _add_mc2i_section_title(self, text: str):
        """
        Titre de section MC2I.
        Point 5 (V5) : couleur corrigée → COLOR_COMPANY RGB(221,0,97)
        au lieu de COLOR_TEXT (gris).
        """
        self._add_section_title(
            text,
            font         = "Lato",
            size         = 12,
            color        = self.COLOR_COMPANY,   # ← corrigé V5 (était COLOR_TEXT)
            bold         = True,
            alignment    = WD_ALIGN_PARAGRAPH.LEFT,
            space_before = 12,
            space_after  = 6,
        )
