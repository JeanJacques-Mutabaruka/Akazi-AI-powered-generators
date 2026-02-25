"""
AKAZI CV Generator V3 — Refactorisé
Hérite de BaseGenerator (Template Method pattern).

Changements vs V2 :
  ✅ Hérite BaseGenerator → supprime load_data, save_document, generate dupliqués
  ✅ setup_document()   → surcharge légère (marges 1.27cm, police 9pt)
  ✅ validate_data()    → signature harmonisée (self.data déjà chargé)
  ✅ generate_content() → remplace generate() comme point d'entrée contenu
  ✅ _make_run()        → remplace les 21 blocs run.font.name/size/color.rgb
  ✅ _apply_para_format()  → remplace _apply_standard_paragraph_format()
  ✅ _add_bullet_item() → hérité de BaseGenerator
  ✅ _add_section_title() → hérité de BaseGenerator
"""

from pathlib import Path
from typing import Dict, Any, List, Optional

from docx.shared import Pt, Cm, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generators.base_generator import BaseGenerator


class AkaziCVGenerator(BaseGenerator):
    """
    Générateur CV au format AKAZI V1.
    Structure JSON attendue : document_metadata, header, skills_table, experience_table
    """

    DEFAULT_FONT = "Century Gothic"
    DEFAULT_SIZE = 9   # Corps de texte AKAZI CV = 9pt

    def __init__(self, input_file: Path, output_file: Path, **kwargs):
        super().__init__(input_file, output_file, **kwargs)
        self.left_col_width:  Optional[Length] = None
        self.right_col_width: Optional[Length] = None

    # =========================================================================
    # SETUP DOCUMENT
    # =========================================================================

    def setup_document(self):
        """Surcharge : marges AKAZI CV (1.27 cm) + largeurs colonnes dynamiques"""
        super().setup_document()

        # Police Normal spécifique AKAZI CV
        normal = self.doc.styles['Normal']
        normal.font.name = self.DEFAULT_FONT
        normal.font.size = Pt(self.DEFAULT_SIZE)

        # Marges AKAZI CV
        self._set_margins(Cm(1.27))

        # Calcul largeurs colonnes tableau 21% / 79%
        self._calculate_dynamic_widths()

    def _calculate_dynamic_widths(self):
        """Calcule dynamiquement les largeurs colonnes 21% / 79% en EMU"""
        section    = self.doc.sections[0]
        usable     = section.page_width - section.left_margin - section.right_margin
        usable_emu = int(usable.emu) if hasattr(usable, 'emu') else int(usable)
        self.left_col_width  = Length(int(usable_emu * 0.21))
        self.right_col_width = Length(int(usable_emu * 0.79))

    def _set_fixed_table_layout(self, table, col_widths):
        """Force le layout fixe et applique les largeurs colonnes + cellules"""
        table.autofit = False
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        table._tbl.tblPr.append(tblLayout)
        for i, w in enumerate(col_widths):
            table.columns[i].width = w
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    # =========================================================================
    # VALIDATE DATA
    # =========================================================================

    def validate_data(self) -> bool:
        required = ['document_metadata', 'header', 'skills_table', 'experience_table']
        missing  = [k for k in required if k not in self.data]
        if missing:
            raise ValueError(f"Clés manquantes dans les données AKAZI CV : {missing}")
        fmt = self.data['document_metadata'].get('format_code')
        if fmt != 'AKAZI_V1':
            raise ValueError(f"format_code incorrect : '{fmt}' (attendu : 'AKAZI_V1')")
        return True

    # =========================================================================
    # GENERATE CONTENT
    # =========================================================================

    def generate_content(self):
        """Pipeline contenu AKAZI CV : header → skills → experiences"""
        self.generate_header()
        self.generate_skills_table()
        self.generate_experiences()

    # =========================================================================
    # HEADER
    # =========================================================================

    def generate_header(self):
        """Ligne 1 : Initiales - Titre - Expérience (ROUGE) | Ligne 2 : TJM + Contact"""
        h = self.data['header']

        # Ligne 1 — ROUGE 11pt centré gras
        l1 = self.doc.add_paragraph()
        l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._make_run(
            l1,
            f"{h['initials']} - {h['title']} - {h['years_of_experience']}",
            size=11, color=self.COLOR_RED, bold=True
        )

        # Ligne 2 — BLEU 11pt + OR pour email
        l2 = self.doc.add_paragraph()
        l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tjm = h.get('daily_rate', '----------')
        self._make_run(l2, f"TJM souhaité : {tjm}€ - Contact : Janvier au 0783388802",
                       size=11, color=self.COLOR_BLUE, bold=True)
        self._make_run(l2, " ou sur ", size=11, color=self.COLOR_BLUE, bold=True)
        self._make_run(l2, "contact@akazi.fr", size=11, color=self.COLOR_GOLD, bold=True)

    # =========================================================================
    # SKILLS TABLE
    # =========================================================================

    def generate_skills_table(self):
        """Tableau compétences 5 lignes × 2 colonnes"""
        s = self.data['skills_table']
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        self._set_fixed_table_layout(table, [self.left_col_width, self.right_col_width])

        self._populate_skill_row(table.rows[0], "COMPÉTENCES\nFONCTIONNELLES",
                                  s['functional_skills']['summary'],
                                  s['functional_skills']['details'])
        self._populate_skill_row(table.rows[1], "COMPÉTENCES\nTECHNIQUES",
                                  s['technical_skills']['summary'],
                                  s['technical_skills']['details'])
        self._populate_education_row(table.rows[2],      s.get('education', []))
        self._populate_certifications_row(table.rows[3], s.get('certifications', []))
        self._populate_languages_row(table.rows[4],      s.get('languages', []))

    def _make_left_cell_title(self, cell, title: str):
        """Cellule gauche : titre centré ROUGE 9pt gras"""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._make_run(para, title, size=9, color=self.COLOR_RED, bold=True)
        cell.vertical_alignment = 1

    def _populate_skill_row(self, row, title: str, summary: str, details: List[Dict]):
        self._make_left_cell_title(row.cells[0], title)
        right = row.cells[1]; right.text = ''
        if summary:
            p = right.add_paragraph()
            self._apply_para_format(p, 3, 3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            self._make_run(p, summary, size=9, color=self.COLOR_BLUE, bold=True)
        for d in details:
            self._add_bullet_item(right, d.get('text', ''),
                                   bold_prefix=d.get('bold_prefix', ''), size=9)

    def _populate_education_row(self, row, education: List[Dict]):
        self._make_left_cell_title(row.cells[0], "FORMATION &\nDIPLOMES")
        right = row.cells[1]; right.text = ''
        for e in education:
            self._add_bullet_item(
                right,
                f"{e.get('year','')} : {e.get('degree','')} - {e.get('institution','')}",
                size=9
            )

    def _populate_certifications_row(self, row, certifications: List):
        self._make_left_cell_title(row.cells[0], "CERTIFICATIONS")
        right = row.cells[1]; right.text = ''
        for c in certifications:
            text = c.get('text', str(c)) if isinstance(c, dict) else str(c)
            self._add_bullet_item(right, text, size=9)

    def _populate_languages_row(self, row, languages: List[Dict]):
        self._make_left_cell_title(row.cells[0], "LANGUES")
        right = row.cells[1]; right.text = ''
        for lang in languages:
            color = self.COLOR_ORANGE if lang.get('inferred') else self.COLOR_BLACK
            self._add_bullet_item(
                right,
                f"{lang.get('language','')} : {lang.get('level','')}",
                color=color, size=9
            )

    # =========================================================================
    # EXPERIENCES
    # =========================================================================

    def generate_experiences(self):
        """Section Expérience Professionnelle"""
        experiences = self.data.get('experience_table', [])
        if not experiences:
            return
        self._add_experience_section_title()
        for exp in experiences:
            self._generate_single_experience(exp)

    def _add_experience_section_title(self):
        """Titre de section dans un tableau pleine largeur centré"""
        self.doc.add_paragraph()
        ttable = self.doc.add_table(rows=1, cols=1)
        ttable.autofit = False
        ttable.style   = 'Table Grid'
        para = ttable.rows[0].cells[0].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._make_run(para, "EXPERIENCE PROFESSIONNELLE",
                       size=11, color=self.COLOR_RED, bold=True)

    def _extract_period(self, period) -> str:
        """Gère period en tant que str OU dict {'formatted': ...}"""
        if isinstance(period, dict):
            return period.get('formatted', 'N/A')
        return str(period) if period else 'N/A'
    
    def _generate_single_experience(self, exp: Dict):
        """Tableau 2 colonnes pour une expérience individuelle"""
        table = self.doc.add_table(rows=1, cols=2)
        table.style   = 'Table Grid'
        table.autofit = False

        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        table._tbl.tblPr.append(tblLayout)
        self._set_fixed_table_layout(table, [self.left_col_width, self.right_col_width])

        self._populate_left_cell(
            table.rows[0].cells[0],
            company=exp.get('company', 'N/A'),
            # period=exp.get('period', {}).get('formatted', 'N/A')
            period=self._extract_period(exp.get('period', 'N/A'))
        )
        self._populate_right_cell(
            table.rows[0].cells[1],
            title   = exp.get('title', 'N/A'),
            context = exp.get('mission_context', ''),
            tasks   = exp.get('tasks', []),
            tech_env= exp.get('technical_environment', [])
        )

    def _populate_left_cell(self, cell, company: str, period: str):
        """Colonne gauche (21%) : Entreprise + Période en BLEU"""
        cell.text = ''; cell.vertical_alignment = 0

        p = cell.add_paragraph()
        self._apply_para_format(p, 3, 3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._make_run(p, company, size=9, color=self.COLOR_BLUE, bold=True)

        cell.add_paragraph()

        p2 = cell.add_paragraph()
        self._apply_para_format(p2, 3, 3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._make_run(p2, f"({period})", size=9, color=self.COLOR_BLUE, bold=True)

    def _populate_right_cell(self, cell, title: str, context: str,
                              tasks: List[str], tech_env: List[Dict]):
        """Colonne droite (79%) : Titre + Contexte + Tâches + Tech"""
        cell.text = ''

        # 1. TITRE
        pt = cell.add_paragraph()
        self._apply_para_format(pt, 6, 6)
        self._make_run(pt, title, size=10, color=self.COLOR_BLUE, bold=True)

        # 2. CONTEXTE — label
        pc = cell.add_paragraph()
        self._apply_para_format(pc, 3, 3, left_indent=0.5,
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        self._make_run(pc, "■ Contexte de la mission : ",
                       size=9, color=self.COLOR_BLACK, bold=True)
        # CONTEXTE — texte
        pct = cell.add_paragraph()
        self._apply_para_format(pct, 6, 6, left_indent=1.0, first_line=0.0)
        pct.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._make_run(pct, context, size=9, color=self.COLOR_BLACK, italic=True)

        # 3. TÂCHES — label
        ptl = cell.add_paragraph()
        self._apply_para_format(ptl, 3, 3, left_indent=0.5,
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        self._make_run(ptl, "■ Tâches :", size=9, color=self.COLOR_BLACK, bold=True)
        # TÂCHES — items
        for task in tasks:
            self._add_bullet_item(cell, task, left_indent=1.5, size=9)

        cell.add_paragraph()

        # 4. TECH — label
        ptel = cell.add_paragraph()
        self._apply_para_format(ptel, 3, 3, left_indent=0.5,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        self._make_run(ptel, "■ Environnements / Outils et technologies : ",
                       size=9, color=self.COLOR_BLACK, bold=True)
        # TECH — texte
        tech_str = ', '.join(filter(None, [
            t.get('name', t.get('technologies', '')) if isinstance(t, dict) else str(t)
            for t in tech_env
        ]))
        ptet = cell.add_paragraph()
        self._apply_para_format(ptet, 6, 6, left_indent=1.0, first_line=0.0)
        ptet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._make_run(ptet, tech_str, size=9, color=self.COLOR_BLACK, bold=True)
