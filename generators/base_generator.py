"""
Base Generator V3 — Séparation contenu / header-footer
Pattern Template Method : generate() orchestre, sous-classes implémentent generate_content().

CHANGEMENTS V3 :
  ✅ setup_document() : ne touche PLUS au header/footer directement
  ✅ _apply_header_footer() : nouveau point d'entrée délégué au HeaderFooterEngine
  ✅ Paramètre hf_preset (Path ou str) accepté dans __init__
  ✅ Rétrocompatibilité : si hf_preset=None → fallback sur AkaziLayoutManager (comportement V2)
"""

from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, Union
from pathlib import Path
from io import BytesIO
import json
import yaml

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utils.logger import logger
from utils.validator import DocumentValidator
from utils.performance import PerformanceTracker


class BaseGenerator(ABC):
    """
    Classe mère abstraite pour tous les générateurs AKAZI.

    Pattern Template Method :
      generate() orchestre → sous-classes implémentent generate_content().

    Nouveau paramètre : hf_preset
      - Path vers un fichier YAML de configuration header/footer
      - Si None → fallback sur le layout corporate AKAZI par défaut
      - "none" (str) → pas de header/footer du tout

    Exemple d'utilisation :
        generator = AkaziCVGenerator(
            input_file=Path("cv.json"),
            output_file=Path("output/cv.docx"),
            hf_preset=Path("hf_presets/headers/akazi_standard.yaml")
        )
    """

    # ── Couleurs corporate AKAZI ──────────────────────────────────────────────
    COLOR_RED     = RGBColor(192,   0,   0)
    COLOR_BLUE    = RGBColor(  0,  32,  96)
    COLOR_GOLD    = RGBColor(204, 153,   0)
    COLOR_ORANGE  = RGBColor(255, 140,   0)
    COLOR_BLACK   = RGBColor(  0,   0,   0)
    # ── Couleurs spécifiques MC2I ─────────────────────────────────────────────
    COLOR_COMPANY = RGBColor(221,   0,  97)
    COLOR_MISSION = RGBColor(  0, 106, 158)
    COLOR_TEXT    = RGBColor( 87,  88,  86)

    DEFAULT_FONT = "Century Gothic"
    DEFAULT_SIZE = 10

    def __init__(
        self,
        input_file:  Path,
        output_file: Path,
        lang:        Optional[str]              = None,
        hf_preset:   Optional[Union[Path, str]] = None,
        **kwargs
    ):
        self.input_file  = Path(input_file)
        self.output_file = Path(output_file)
        self.lang        = lang or 'en'
        self.data: Optional[Dict[str, Any]] = None
        self.doc:  Optional[Document]       = None

        # ── Preset header/footer ──────────────────────────────────────────────
        # None          → layout corporate par défaut (rétrocompatibilité)
        # "none"        → pas de header/footer
        # Path/str      → chemin vers un fichier YAML de preset
        if hf_preset is not None and str(hf_preset).lower() != "none":
            self.hf_preset = Path(hf_preset)
        elif str(hf_preset).lower() == "none" if hf_preset else False:
            self.hf_preset = "none"
        else:
            self.hf_preset = None  # fallback comportement V2

        self.validator           = DocumentValidator()
        self.performance_tracker = PerformanceTracker()

        if not self.input_file.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_file}")

        logger.debug(
            "generator_initialized",
            generator=self.__class__.__name__,
            input_file=str(self.input_file),
            output_file=str(self.output_file),
            hf_preset=str(self.hf_preset) if self.hf_preset else "default"
        )

    # =========================================================================
    # TEMPLATE METHOD — ne pas surcharger dans les sous-classes
    # =========================================================================

    def generate(self) -> Path:
        """
        Pipeline principal :
          load_data → validate_data → setup_document
          → _apply_header_footer → generate_content → save_document
        """
        try:
            logger.info(
                "generation_started",
                generator=self.__class__.__name__,
                input_file=self.input_file.name
            )
            if self.data is None:
                self.load_data()
            self.validate_data()
            self.setup_document()
            self._apply_header_footer()   # ← NOUVEAU : séparé de setup_document
            self.generate_content()
            self.save_document()
            logger.info(
                "generation_completed",
                generator=self.__class__.__name__,
                output_file=self.output_file.name
            )
            return self.output_file

        except Exception as e:
            logger.error(
                "generation_failed",
                generator=self.__class__.__name__,
                error=str(e)
            )
            raise

    def generate_to_buffer(self) -> BytesIO:
        """Génère le document dans un buffer BytesIO (usage Streamlit)"""
        self.generate()
        buf = BytesIO()
        with open(self.output_file, 'rb') as f:
            buf.write(f.read())
        buf.seek(0)
        return buf

    # =========================================================================
    # ÉTAPES COMMUNES
    # =========================================================================

    def load_data(self) -> Dict[str, Any]:
        """Charge JSON ou YAML"""
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                suffix = self.input_file.suffix.lower()
                if suffix == '.json':
                    self.data = json.load(f)
                elif suffix in ('.yaml', '.yml'):
                    self.data = yaml.safe_load(f)
                else:
                    raise ValueError(f"Format non supporté : {suffix}")

            if not isinstance(self.data, dict):
                raise ValueError("Les données doivent être un dictionnaire")
            return self.data

        except (json.JSONDecodeError, yaml.YAMLError) as e:
            raise ValueError(f"Erreur de parsing : {e}")

    def save_document(self):
        """Sauvegarde le document Word"""
        self.output_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_file))
        logger.info("document_saved", output=str(self.output_file))

    def setup_document(self):
        """
        Initialise le document Word avec styles AKAZI.
        NE CONFIGURE PLUS le header/footer — délégué à _apply_header_footer().
        Les sous-classes peuvent surcharger pour ajuster police/marges.
        """
        self.doc = Document()

        # Styles centralisés
        try:
            from utils.akazi_styles import AkaziStyleManager
            AkaziStyleManager(self.doc).setup_all_styles()
        except Exception as e:
            logger.warning("style_manager_skipped", reason=str(e))

        # Police et marges par défaut
        normal = self.doc.styles['Normal']
        normal.font.name = self.DEFAULT_FONT
        normal.font.size = Pt(self.DEFAULT_SIZE)
        self._set_margins(Cm(1.27))

    def _apply_header_footer(self):
        """
        Applique le header/footer selon la stratégie choisie :

        Stratégie 1 — Preset YAML fourni :
            Utilise HeaderFooterEngine.apply_yaml() (moteur générique)

        Stratégie 2 — Pas de preset (None) :
            Fallback sur AkaziLayoutManager (comportement V2, rétrocompatibilité)

        Stratégie 3 — hf_preset == "none" :
            Aucun header/footer appliqué

        Les sous-classes peuvent surcharger cette méthode pour un contrôle total
        (ex: MC2ICVGenerator qui utilise MC2ILayoutManager avec données dynamiques).
        """
        # Stratégie 3 : pas de H/F
        if self.hf_preset == "none":
            logger.debug("hf_skipped", reason="hf_preset='none'")
            return

        # Stratégie 1 : preset YAML
        if self.hf_preset is not None:
            if not Path(self.hf_preset).exists():
                logger.warning(
                    "hf_preset_not_found",
                    path=str(self.hf_preset),
                    fallback="default_akazi_layout"
                )
                self._apply_default_header_footer()
                return

            try:
                from docx_header_engine.engine import HeaderFooterEngine
                engine = HeaderFooterEngine(self.doc)
                engine.apply_yaml(str(self.hf_preset))
                logger.info("hf_preset_applied", preset=str(self.hf_preset))
                return
            except Exception as e:
                import traceback as _tb
                _detail = _tb.format_exc()
                logger.error("hf_preset_apply_failed", error=str(e), detail=_detail)
                # Stocker l'erreur pour remontée dans Streamlit
                self._hf_error = (
                    f"{type(e).__name__}: {e}\n"
                    f"Preset: {self.hf_preset}\n"
                    f"{_detail}"
                )
                # Fallback sur défaut (le document est quand même généré)
                self._apply_default_header_footer()
                return

        # Stratégie 2 : fallback V2
        self._apply_default_header_footer()

    def _apply_default_header_footer(self):
        """Applique le layout AKAZI corporate par défaut (comportement V2)"""
        try:
            from utils.akazi_layout import AkaziLayoutManager
            AkaziLayoutManager(self.doc).setup_header_footer()
            logger.debug("default_akazi_layout_applied")
        except Exception as e:
            logger.warning("layout_manager_skipped", reason=str(e))

    def _set_margins(self, margin: Cm):
        """Applique des marges uniformes à toutes les sections"""
        for section in self.doc.sections:
            section.top_margin    = margin
            section.bottom_margin = margin
            section.left_margin   = margin
            section.right_margin  = margin

    # =========================================================================
    # MÉTHODES ABSTRAITES
    # =========================================================================

    @abstractmethod
    def validate_data(self) -> bool:
        pass

    @abstractmethod
    def generate_content(self):
        pass

    # =========================================================================
    # HELPERS COMMUNS
    # =========================================================================

    def _make_run(self, para, text, font=None, size=None, color=None,
                  bold=False, italic=False, small_caps=False):
        run = para.add_run(text)
        run.font.name       = font  or self.DEFAULT_FONT
        run.font.size       = Pt(size or self.DEFAULT_SIZE)
        run.font.color.rgb  = color or self.COLOR_BLACK
        run.bold            = bold
        run.italic          = italic
        run.font.small_caps = small_caps
        return run

    def _apply_para_format(self, para, space_before=3, space_after=3,
                            left_indent=None, first_line=None,
                            alignment=None, line_spacing=None):
        pf = para.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        if left_indent  is not None: pf.left_indent       = Cm(left_indent)
        if first_line   is not None: pf.first_line_indent = Cm(first_line)
        if alignment    is not None: para.alignment        = alignment
        if line_spacing is not None: pf.line_spacing       = line_spacing

    def _add_section_title(self, text, font=None, size=12, color=None,
                            bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_before=12, space_after=6, container=None):
        target = container if container is not None else self.doc
        para   = target.add_paragraph()
        self._apply_para_format(para, space_before, space_after, alignment=alignment)
        self._make_run(para, text,
                        font=font or self.DEFAULT_FONT, size=size,
                        color=color or self.COLOR_RED, bold=bold)
        return para

    def _add_bullet_item(self, container, text, bold_prefix='', color=None,
                          left_indent=0.5, hanging=0.5, size=None, font=None,
                          italic=False, level=0):
        if color is None:
            color = self.COLOR_BLACK

        para = container.add_paragraph()
        style_name = f'Akazi Bullet Level {min(level + 1, 2)}'
        try:
            para.style = style_name
        except KeyError:
            para.style = 'List Bullet'

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.left_indent       = Cm(left_indent)
        pf.first_line_indent = Cm(-hanging)
        pf.space_before      = Pt(3)
        pf.space_after       = Pt(3)

        try:
            from utils.akazi_styles import AkaziStyleManager
            AkaziStyleManager.apply_bullet(para, level=level)
        except Exception:
            pass

        _font = font or self.DEFAULT_FONT
        _size = size or self.DEFAULT_SIZE

        if bold_prefix:
            br = para.add_run(bold_prefix + " : ")
            br.font.name = _font; br.font.size = Pt(_size)
            br.font.color.rgb = color; br.bold = True
            rest = text.replace(bold_prefix + " : ", "").replace(bold_prefix + ": ", "")
            nr = para.add_run(rest)
            nr.font.name = _font; nr.font.size = Pt(_size)
            nr.font.color.rgb = color; nr.italic = italic
        else:
            r = para.add_run(text)
            r.font.name = _font; r.font.size = Pt(_size)
            r.font.color.rgb = color; r.italic = italic

        return para

    def _add_horizontal_separator(self, color_hex="DD0061"):
        para = self.doc.add_paragraph()
        pPr  = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return para

    def _disable_contextual_spacing(self, style_name):
        try:
            style = self.doc.styles[style_name]
            pPr   = style.element.get_or_add_pPr()
            contextual = pPr.find(qn('w:contextualSpacing'))
            if contextual is not None:
                pPr.remove(contextual)
        except KeyError:
            pass
