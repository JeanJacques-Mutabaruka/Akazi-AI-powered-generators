"""
mc2i_layout.py
----------------------------------------
Gestion corporate du Header & Footer MC2I
Positionnement absolu des images via XML direct

Header MC2I :
  - Logo MC2I (droite) : 4.6cm × 2.88cm, position absolue 14cm right, -0.8cm below
  - Texte (gauche) 2 lignes : Trigram / Titre - Années, Lato 16pt, RGB(0,106,158)

Footer MC2I :
  - Image points (droite) : 3.37cm × 7.96cm, position absolue 15.25cm right, -6.5cm below
  - Contact (gauche) 2 lignes : Nom-Email / Adresse, Lato 10pt, RGB(0,106,158)
  - Pagination (droite) : Page X/Y, Lato 10pt, RGB(0,106,158)
"""

from pathlib import Path

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class MC2ILayoutManager:
    """
    Gestionnaire de Header/Footer pour documents MC2I.
    
    Données consultant extraites du JSON MC2I :
      - Trigram : data['document_metadata']['consultant_initials']
      - Titre : data['professional_experiences'][0]['title']
      - Années : calculées depuis data['experience_summary']
    """

    MC2I_COLOR = RGBColor(0, 106, 158)  # Bleu MC2I corporate

    def __init__(self, document, data=None, project_root=None):
        """
        Args:
            document: Document python-docx
            data: Dictionnaire JSON MC2I complet
            project_root: Racine du projet (pour trouver Assets/)
        """
        self.doc = document
        self.data = data or {}
        
        if project_root:
            self.project_root = Path(project_root)
        else:
            self.project_root = Path(__file__).resolve().parent.parent

    # ==========================================================
    # PUBLIC
    # ==========================================================

    def setup_header_footer(self):
        """Configure header et footer MC2I"""
        self._setup_header()
        self._setup_footer()

    # ==========================================================
    # HEADER
    # ==========================================================

    def _setup_header(self):
        """
        Header MC2I :
          - Logo MC2I (droite, position absolue)
          - Texte 2 lignes (gauche) :
            Ligne 1 : Trigram (initiales 3 caractères)
            Ligne 2 : Titre - (Années d'expérience)
        """
        section = self.doc.sections[0]
        header  = section.header
        header.is_linked_to_previous = False

        # Distance from top
        section.header_distance = Cm(0.5)

        para = header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ---- EXTRAIRE DONNÉES DU JSON ----
        # Trigram depuis document_metadata
        metadata = self.data.get('document_metadata', {})
        trigram  = metadata.get('consultant_initials', 'XXX')

        # Titre depuis première expérience professionnelle
        experiences = self.data.get('professional_experiences', [])
        title = experiences[0].get('title', 'Consultant') if experiences else 'Consultant'

        # Années calculées depuis experience_summary
        exp_summary   = self.data.get('experience_summary', [])
        total_months  = sum(e.get('duration_months', 0) for e in exp_summary)
        total_years   = total_months // 12
        years_text    = f"{total_years} années d'expérience"

        # ---- TEXTE HEADER (2 lignes) ----
        # Ligne 1 : Trigram
        run1 = para.add_run(trigram)
        run1.font.name      = "Lato"
        run1.font.size      = Pt(16)
        run1.font.color.rgb = self.MC2I_COLOR
        run1.bold           = True
        para.add_run('\n')

        # Ligne 2 : Titre - (Années)
        run2 = para.add_run(f"{title} – ({years_text})")
        run2.font.name      = "Lato"
        run2.font.size      = Pt(16)
        run2.font.color.rgb = self.MC2I_COLOR
        run2.bold           = False  # Pas bold pour ligne 2

        # ---- LOGO MC2I (position absolue) ----
        logo_path = self.project_root / "Assets" / "images" / "MC2i_logo_header.jpg"
        
        if logo_path.exists():
            self._add_positioned_image(
                para,
                str(logo_path),
                width_cm=4.6,
                height_cm=2.88,
                pos_h_cm=14.0,    # 14 cm to the right of column
                pos_v_cm=-0.8,    # -0.8 cm below paragraph
            )

    # ==========================================================
    # FOOTER
    # ==========================================================

    def _setup_footer(self):
        """
        Footer MC2I :
          - Image points (droite, position absolue)
          - Contact 2 lignes (gauche) : BENJAMIN DUPUY / Adresse
          - Pagination (droite) : Page X/Y
        """
        section = self.doc.sections[0]
        footer  = section.footer
        footer.is_linked_to_previous = False

        section.footer_distance = Cm(1.0)

        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ---- TEXTE CONTACT (2 lignes FIXES) ----
        # Ligne 1 : Nom - Email
        run1 = para.add_run("BENJAMIN DUPUY – benjamin.dupuy@mc2i.fr")
        run1.font.name      = "Lato"
        run1.font.size      = Pt(10)
        run1.font.color.rgb = self.MC2I_COLOR
        run1.bold           = False
        para.add_run('\n')

        # Ligne 2 : Adresse
        run2 = para.add_run("51 Rue François 1er – 75008 PARIS - 01.44.43.01.00 – www.mc2i.fr")
        run2.font.name      = "Lato"
        run2.font.size      = Pt(10)
        run2.font.color.rgb = self.MC2I_COLOR
        run2.bold           = False

        # ---- PAGINATION (côté droit via tab stop) ----
        # Calculer largeur utile
        page_width   = section.page_width
        left_margin  = section.left_margin
        right_margin = section.right_margin
        usable_width = page_width - left_margin - right_margin

        # Tab stop aligné à droite
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

        # Ajouter tab + pagination
        para.add_run("\t")
        
        # Page courante
        page_run = para.add_run()
        page_run.font.name      = "Lato"
        page_run.font.size      = Pt(10)
        page_run.font.color.rgb = self.MC2I_COLOR
        page_run.bold           = False
        self._add_field(page_run, "PAGE")
        
        # Séparateur
        sep_run = para.add_run("/")
        sep_run.font.name      = "Lato"
        sep_run.font.size      = Pt(10)
        sep_run.font.color.rgb = self.MC2I_COLOR
        sep_run.bold           = False
        
        # Total pages
        total_run = para.add_run()
        total_run.font.name      = "Lato"
        total_run.font.size      = Pt(10)
        total_run.font.color.rgb = self.MC2I_COLOR
        total_run.bold           = False
        self._add_field(total_run, "NUMPAGES")

        # ---- IMAGE FOOTER (position absolue) ----
        image_path = self.project_root / "Assets" / "images" / "MC2i_image_footer.jpg"
        
        if image_path.exists():
            self._add_positioned_image(
                para,
                str(image_path),
                width_cm=3.37,
                height_cm=7.96,
                pos_h_cm=15.25,   # 15.25 cm to the right of column
                pos_v_cm=-6.5,    # -6.5 cm below paragraph
            )

    # ==========================================================
    # UTILITIES
    # ==========================================================

    def _add_field(self, run, field_code: str):
        """Ajoute un champ Word (PAGE, NUMPAGES, etc.)"""
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = field_code
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

    def _add_positioned_image(
        self,
        para,
        image_path: str,
        width_cm: float,
        height_cm: float,
        pos_h_cm: float,
        pos_v_cm: float,
    ):
        """
        Ajoute une image avec positionnement absolu via XML direct.
        
        Args:
            para: Paragraphe où insérer l'image
            image_path: Chemin de l'image
            width_cm: Largeur en cm
            height_cm: Hauteur en cm
            pos_h_cm: Position horizontale absolue (cm à droite de la colonne)
            pos_v_cm: Position verticale absolue (cm en dessous du paragraphe)
        """
        # Ajouter l'image normalement d'abord
        run = para.add_run()
        run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))

        # Récupérer l'élément drawing
        drawing = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if not drawing:
            return  # Image pas trouvée, abandon

        # Modifier inline → anchor pour positionnement absolu
        inline = drawing[0].find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if inline is None:
            return

        # Convertir inline en anchor
        anchor = self._convert_inline_to_anchor(
            inline,
            width_cm,
            height_cm,
            pos_h_cm,
            pos_v_cm
        )
        
        # Remplacer inline par anchor
        parent = inline.getparent()
        parent.remove(inline)
        parent.append(anchor)

    def _convert_inline_to_anchor(
        self,
        inline,
        width_cm: float,
        height_cm: float,
        pos_h_cm: float,
        pos_v_cm: float,
    ):
        """
        Convertit un <wp:inline> en <wp:anchor> avec positionnement absolu.
        Utilise les spécifications exactes du PDF MC2I.
        """
        wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        
        # Créer anchor
        anchor = OxmlElement(f'{{{wp_ns}}}anchor')
        
        # Attributs anchor
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251658240')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')

        # simplePos (ignoré mais requis)
        simplePos = OxmlElement(f'{{{wp_ns}}}simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)

        # positionH (horizontal) — "to the right of column"
        positionH = OxmlElement(f'{{{wp_ns}}}positionH')
        positionH.set('relativeFrom', 'column')
        posOffset = OxmlElement(f'{{{wp_ns}}}posOffset')
        posOffset.text = str(int(pos_h_cm * 360000))  # cm → EMU (1cm = 360000 EMU)
        positionH.append(posOffset)
        anchor.append(positionH)

        # positionV (vertical) — "below paragraph"
        positionV = OxmlElement(f'{{{wp_ns}}}positionV')
        positionV.set('relativeFrom', 'paragraph')
        posOffset = OxmlElement(f'{{{wp_ns}}}posOffset')
        posOffset.text = str(int(pos_v_cm * 360000))  # cm → EMU
        positionV.append(posOffset)
        anchor.append(positionV)

        # extent (taille)
        extent = OxmlElement(f'{{{wp_ns}}}extent')
        extent.set('cx', str(int(width_cm * 360000)))
        extent.set('cy', str(int(height_cm * 360000)))
        anchor.append(extent)

        # Wrapping (square)
        wrapSquare = OxmlElement(f'{{{wp_ns}}}wrapSquare')
        wrapSquare.set('wrapText', 'bothSides')
        anchor.append(wrapSquare)

        # Copier le contenu graphique depuis inline
        for child in inline:
            if 'graphic' in child.tag.lower():
                anchor.append(child)

        return anchor
