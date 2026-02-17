"""
akazi_layout.py
----------------------------------------
Gestion corporate du Header & Footer AKAZI
Version ajustée dimensions & corporate color
"""

from pathlib import Path

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class AkaziLayoutManager:

    CORPORATE_COLOR = RGBColor(0, 32, 96)

    def __init__(self, document, project_root=None):
        self.doc = document

        if project_root:
            self.project_root = Path(project_root)
        else:
            self.project_root = Path(__file__).resolve().parent.parent

    # ==========================================================
    # PUBLIC
    # ==========================================================

    def setup_header_footer(self):
        self._setup_header()
        self._setup_footer()

    # ==========================================================
    # HEADER
    # ==========================================================

    def _setup_header(self):
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # 🔹 Distance from top = 0.15 cm
        section.header_distance = Cm(0.15)

        para = header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- LOGO ----
        logo_path = self.project_root / "Assets" / "images" / "Akazi_logo_small.jpg"

        if logo_path.exists():
            run = para.add_run()
            run.add_picture(
                str(logo_path),
                width=Cm(2.7),
                height=Cm(1.4)
            )
            run.add_break()

        # ---- TEXT ----
        run = para.add_run("Vos consultants, de A à Z")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = self.CORPORATE_COLOR

    # ==========================================================
    # FOOTER
    # ==========================================================


    def _setup_footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Distance from bottom = 0.95 cm
        section.footer_distance = Cm(0.95)

        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Important

        # ======================================================
        # CALCUL LARGEUR UTILE A4
        # ======================================================

        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin

        usable_width = page_width - left_margin - right_margin

        # ======================================================
        # TAB STOP DROITE (position exacte)
        # ======================================================

        tab_stops = para.paragraph_format.tab_stops
        tab_stops.clear_all()  # sécurité

        tab_stops.add_tab_stop(
            usable_width,
            WD_TAB_ALIGNMENT.RIGHT
        )

        # ======================================================
        # TEXTE SOCIETE (CENTRÉ MANUELLEMENT)
        # ======================================================

        company_text = "Société AKAZI, SAS – 60 Rue François 1er 75008 Paris"

        run = para.add_run(company_text)
        self._format_footer_run(run)

        # Centrage visuel réel
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # ======================================================
        # TAB → pagination alignée droite
        # ======================================================

        para.add_run("\t")

        page_run = para.add_run("Page ")
        self._format_footer_run(page_run)
        self._add_field(page_run, "PAGE")

        de_run = para.add_run(" de ")
        self._format_footer_run(de_run)

        total_run = para.add_run()
        self._format_footer_run(total_run)
        self._add_field(total_run, "NUMPAGES")

    # ==========================================================
    # UTILITIES
    # ==========================================================

    def _format_footer_run(self, run):
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.bold = True
        run.font.color.rgb = self.CORPORATE_COLOR

    def _add_field(self, run, field_code):
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = field_code
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)
