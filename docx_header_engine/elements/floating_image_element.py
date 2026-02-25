import os
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from .base_element import BaseElement


class FloatingImageElement(BaseElement):
    """
    Image flottante positionnée à des coordonnées absolues (OpenXML anchor).

    COMPORTEMENT :
    - L'anchor est toujours injecté directement dans le body du header/footer
      (jamais dans une cellule de table), ce qui est requis par Word pour que
      behindDoc et la position absolue fonctionnent correctement.
    - Le paragraphe de la cellule reste vide (carrier neutre).

    Config:
        path        : chemin vers l'image
        width_cm    : largeur en cm
        height_cm   : hauteur en cm (optionnel → proportionnel)
        x_cm        : distance depuis le bord GAUCHE de la page
        y_cm        : distance depuis le HAUT de la page
                      ⚠ Pour un footer standard A4 (marges 2.54cm) :
                        zone footer = 25.4cm → 26.7cm depuis le haut
                        → utiliser y_cm entre 25.4 et 26.0
        z_order     : ordre Z (entier >= 1, défaut 1)
        wrap        : "none" | "tight" | "square" (défaut "none")
        behind_text : True → image derrière le texte (défaut False)

    Calcul x_cm pour aligner le bord DROIT :
        x_cm = position_bord_droit_cible - width_cm
        ex : bord droit à 19cm, image 3.5cm → x_cm = 15.5

    Helper statique disponible :
        FloatingImageElement.footer_y(section) → y_cm optimal pour footer
        FloatingImageElement.header_y(section) → y_cm optimal pour header
    """

    EMU_PER_CM = 360000
    NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    NS_A  = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    def __init__(self, config):
        self.path        = config.get("path")
        self.width_cm    = config.get("width_cm", 3.0)
        self.height_cm   = config.get("height_cm")
        self.x_cm        = config.get("x_cm", 0.0)
        self.y_cm        = config.get("y_cm", 0.0)
        self.z_order     = max(1, config.get("z_order", 1))
        self.wrap        = config.get("wrap", "none")
        self.behind_text = config.get("behind_text", False)

    def render(self, container):
        """
        Injecte l'anchor dans le BODY du header/footer (pas dans la cellule).
        container peut être une cellule de table ou directement un header/footer.
        """
        if not self.path or not os.path.exists(self.path):
            raise FileNotFoundError(
                f"FloatingImageElement: image introuvable → {self.path}"
            )

        # Remonter jusqu'au body du header/footer (w:hdr ou w:ftr)
        hf_body = self._find_hf_body(container)

        # Créer un paragraphe temporaire pour que python-docx génère la relation image
        # On l'ajoute dans le container (cellule) puis on extrait l'anchor
        p_tmp = container.add_paragraph()
        pic_run = p_tmp.add_run()
        pic_run.add_picture(
            self.path,
            width=Cm(self.width_cm),
            height=Cm(self.height_cm) if self.height_cm else None
        )

        inline = pic_run._r.find(f'.//{{{self.NS_WP}}}inline')
        if inline is None:
            return

        # Construire l'anchor
        anchor = self._build_anchor(inline)

        # Vider p_tmp mais le CONSERVER dans la cellule :
        # OOXML exige qu'une cellule ait toujours au moins un <w:p>.
        # Le supprimer laisserait la cellule vide → "unreadable content" dans Word.
        for child in list(p_tmp._element):
            p_tmp._element.remove(child)

        # Injecter l'anchor dans le body du header/footer via un paragraphe dédié
        # Ce paragraphe est inséré EN PREMIER dans le hf_body (avant la table)
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        p_anchor = OxmlElement('w:p')
        r_anchor = OxmlElement('w:r')
        drawing  = OxmlElement('w:drawing')
        drawing.append(anchor)
        r_anchor.append(drawing)
        p_anchor.append(r_anchor)

        # Insérer au début du hf_body (index 0) pour que l'image soit "sous" la table
        hf_body.insert(0, p_anchor)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _find_hf_body(container):
        """
        Remonte la hiérarchie XML pour trouver le nœud w:hdr ou w:ftr.
        Fonctionne que container soit une cellule, un paragraphe ou le footer lui-même.
        """
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        elem = getattr(container, '_element', container)

        # Remonter jusqu'à trouver hdr ou ftr
        node = elem
        while node is not None:
            tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
            if tag in ('hdr', 'ftr'):
                return node
            node = node.getparent()

        # Fallback : retourner l'élément lui-même
        return elem

    @staticmethod
    def footer_y(section):
        """
        Retourne le y_cm optimal pour placer une image dans le footer.
        = distance depuis le haut de la page jusqu'au début de la zone footer.
        """
        EMU = FloatingImageElement.EMU_PER_CM
        return (section.page_height - section.bottom_margin) / EMU

    @staticmethod
    def header_y(section):
        """
        Retourne le y_cm optimal pour placer une image dans le header.
        = header_distance depuis le haut de la page.
        """
        EMU = FloatingImageElement.EMU_PER_CM
        return section.header_distance / EMU

    def _build_anchor(self, inline):
        """
        Convertit un <wp:inline> en <wp:anchor> avec position absolue.
        Ordre des enfants imposé par le schéma OOXML :
          simplePos → positionH → positionV → extent → effectExtent
          → wrap → docPr → cNvGraphicFramePr → a:graphic
        """
        x_emu = int(self.x_cm * self.EMU_PER_CM)
        y_emu = int(self.y_cm * self.EMU_PER_CM)

        extent  = inline.find(f'{{{self.NS_WP}}}extent')
        doc_pr  = inline.find(f'{{{self.NS_WP}}}docPr')
        cnv     = inline.find(f'{{{self.NS_WP}}}cNvGraphicFramePr')
        graphic = inline.find(f'{{{self.NS_A}}}graphic')

        anchor = OxmlElement('wp:anchor')
        anchor.set('distT',          '0')
        anchor.set('distB',          '0')
        anchor.set('distL',          '114300')
        anchor.set('distR',          '114300')
        anchor.set('simplePos',      '0')
        anchor.set('relativeHeight', str(self.z_order))
        anchor.set('behindDoc',      '1' if self.behind_text else '0')
        anchor.set('locked',         '0')
        anchor.set('layoutInCell',   '1')
        anchor.set('allowOverlap',   '1')

        sp = OxmlElement('wp:simplePos')
        sp.set('x', '0'); sp.set('y', '0')
        anchor.append(sp)

        pos_h = OxmlElement('wp:positionH')
        pos_h.set('relativeFrom', 'page')
        off_h = OxmlElement('wp:posOffset')
        off_h.text = str(x_emu)
        pos_h.append(off_h)
        anchor.append(pos_h)

        pos_v = OxmlElement('wp:positionV')
        pos_v.set('relativeFrom', 'page')
        off_v = OxmlElement('wp:posOffset')
        off_v.text = str(y_emu)
        pos_v.append(off_v)
        anchor.append(pos_v)

        if extent is not None:
            anchor.append(copy.deepcopy(extent))

        ee = OxmlElement('wp:effectExtent')
        for attr in ('l', 't', 'r', 'b'):
            ee.set(attr, '0')
        anchor.append(ee)

        wrap_map = {'none': 'wp:wrapNone', 'tight': 'wp:wrapTight', 'square': 'wp:wrapSquare'}
        anchor.append(OxmlElement(wrap_map.get(self.wrap, 'wp:wrapNone')))

        if doc_pr is not None:
            anchor.append(copy.deepcopy(doc_pr))
        if cnv is not None:
            anchor.append(copy.deepcopy(cnv))
        if graphic is not None:
            anchor.append(copy.deepcopy(graphic))
        else:
            raise ValueError(
                f"FloatingImageElement: <a:graphic> introuvable. NS attendu: {self.NS_A}"
            )

        return anchor
