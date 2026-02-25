from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base_element import BaseElement


class InlineGroupElement(BaseElement):
    """
    Groupe d'éléments rendus sur UNE SEULE ligne / paragraphe.
    Permet de faire "Page X / Y" en combinant textes statiques et champs Word
    sans créer plusieurs paragraphes séparés.

    Config example:
        {
            "type": "inline_group",
            "align": "right",
            "style": {"font": "Calibri", "size": 9, "color": "666666"},
            "items": [
                {"type": "text",  "value": "Page "},
                {"type": "field", "value": " PAGE "},
                {"type": "text",  "value": " / "},
                {"type": "field", "value": " NUMPAGES "}
            ]
        }
    """

    def __init__(self, config):
        self.items = config.get("items", [])
        self.style = config.get("style", {})
        self.align = config.get("align", "left")

    def render(self, container):
        p = container.add_paragraph()

        # Alignement du paragraphe
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
            "left":   WD_ALIGN_PARAGRAPH.LEFT,
        }
        p.alignment = align_map.get(self.align, WD_ALIGN_PARAGRAPH.LEFT)

        for item in self.items:
            item_type = item.get("type")
            item_style = {**self.style, **item.get("style", {})}  # merge styles

            if item_type == "text":
                run = p.add_run(item.get("value", ""))
                self._apply_style(run, item_style)

            elif item_type == "field":
                run = p.add_run()
                self._apply_style(run, item_style)
                self._add_field(run, item.get("value", " PAGE "))

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _apply_style(run, style):
        if style.get("font"):
            run.font.name = style["font"]
        if style.get("size"):
            run.font.size = Pt(style["size"])
        if style.get("bold") is not None:
            run.bold = style["bold"]
        if style.get("italic") is not None:
            run.italic = style["italic"]
        if style.get("color"):
            run.font.color.rgb = RGBColor.from_string(style["color"])

    @staticmethod
    def _add_field(run, instruction):
        """Insère un champ Word (fldChar) dans le run fourni."""
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instruction

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)
