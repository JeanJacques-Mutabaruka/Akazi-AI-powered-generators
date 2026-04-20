import warnings
from pathlib import Path
from docx.shared import Cm
from .base_element import BaseElement


class ImageElement(BaseElement):

    def __init__(self, config):
        self.path = config.get("path")
        self.width_cm = config.get("width_cm")
        self.height_cm = config.get("height_cm")

    def _resolve_path(self):
        if not self.path:
            return None
        p = Path(self.path)
        if p.exists():
            return str(p)
        project_root = Path(__file__).resolve().parent.parent.parent
        candidate = project_root / "Assets" / "images" / p.name
        return str(candidate) if candidate.exists() else None

    def render(self, container):
        resolved = self._resolve_path()
        if not resolved:
            container.add_paragraph()
            warnings.warn(f"[ImageElement] Image introuvable, ignorée : {self.path}", stacklevel=2)
            return
        p = container.add_paragraph()

        if self.width_cm and self.height_cm:
            p.add_run().add_picture(
                resolved,
                width=Cm(self.width_cm),
                height=Cm(self.height_cm)
            )

        elif self.width_cm:
            p.add_run().add_picture(
                resolved,
                width=Cm(self.width_cm)
            )

        else:
            p.add_run().add_picture(resolved)
